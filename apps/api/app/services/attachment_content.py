"""Attachment MIME allowlists and text extraction for chat context."""

from __future__ import annotations

import asyncio
import base64
import io
import logging
import zipfile
from dataclasses import dataclass
from typing import Any
from uuid import UUID

from sqlalchemy.ext.asyncio import AsyncSession

from app.core.attachment_limits import MAX_ATTACHMENT_SIZE as MAX_ATTACHMENT_SIZE
from app.core.config import Settings
from app.gateways.storage_gateway import StorageGateway

logger = logging.getLogger(__name__)

# Inline excerpt cap — keeps the user-message attachment excerpt small so it
# doesn't bloat the prompt with a full document dump.
MAX_EXTRACT_CHARS = 12_000
# Indexing cap — the background indexing path chunks the whole document for
# pgvector RAG, so it can extract far more than the inline excerpt. Capped to
# avoid runaway memory on adversarially large text payloads.
MAX_INDEX_EXTRACT_CHARS = 200_000
# Inline excerpts stay small; background indexing can read later PDF pages.
PDF_EXTRACT_MAX_PAGES = 25
PDF_INDEX_MAX_PAGES = 500
VISION_REHYDRATE_LIMIT = 2

_IMAGE_UNAVAILABLE_NOTE = (
    "The earlier image is not available to look at again. "
    "Do not guess from a prior description; say you cannot see it."
)
_ATTACH_ID_PREFIX = "/attachments/"
_ATTACH_ID_SUFFIX = "/file"

GALLERY_THUMB_MIN_EDGE = 32
GALLERY_THUMB_MAX_EDGE = 512


def resize_image_bytes(data: bytes, edge: int) -> tuple[bytes, str]:
    """Downscale an image for gallery tiles. Returns (bytes, content_type)."""
    from PIL import Image

    size = max(GALLERY_THUMB_MIN_EDGE, min(GALLERY_THUMB_MAX_EDGE, edge))
    opened = Image.open(io.BytesIO(data))
    opened.thumbnail((size, size))
    out = io.BytesIO()
    if opened.mode in ("RGBA", "LA", "P"):
        opened.save(out, format="PNG", optimize=True)
        return out.getvalue(), "image/png"
    converted = opened.convert("RGB") if opened.mode != "RGB" else opened
    converted.save(out, format="JPEG", quality=80, optimize=True)
    return out.getvalue(), "image/jpeg"


IMAGE_CONTENT_TYPES = frozenset(
    {
        "image/jpeg",
        "image/jpg",
        "image/png",
        "image/webp",
        "image/gif",
        # HEIC/HEIF rejected — unreliable cross-platform preview + vision; ask for JPEG/PNG.
    }
)

DOCUMENT_CONTENT_TYPES = frozenset(
    {
        "application/pdf",
        "text/plain",
        "text/markdown",
        "text/csv",
        "application/json",
        # Legacy .doc (application/msword) intentionally excluded — no pure-Python
        # parser can read it, so uploads succeed but RAG silently fails. Reject
        # at presign with a clear error so users convert to .docx instead.
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    }
)

ALLOWED_CONTENT_TYPES = IMAGE_CONTENT_TYPES | DOCUMENT_CONTENT_TYPES

_CONTENT_TYPE_ALIASES = {
    "image/jpg": "image/jpeg",
    "image/pjpeg": "image/jpeg",
}

# Text-ish types have no reliable magic signature. Accept only clean UTF-8
# with no NUL (rejects polyglot binaries claimed as txt/json/csv/md). HEIC/HEIF
# use an ISO BMFF ftyp box that's awkward to sniff here and are rare via the
# local upload path, so they're accepted on trust for now.
_TEXTISH_TYPES = frozenset({"text/plain", "text/markdown", "text/csv", "application/json"})
_UNVERIFIABLE_TYPES: frozenset[str] = frozenset()


_DOCX_MARKER_ENTRY = "word/document.xml"
_MAX_DOCX_TOTAL_UNCOMPRESSED = 32 * 1024 * 1024
_OFFICE_MARKERS = {
    _DOCX_MARKER_ENTRY: "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xl/workbook.xml": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "ppt/presentation.xml": (
        "application/vnd.openxmlformats-officedocument.presentationml.presentation"
    ),
}


def _office_content_type(data: bytes) -> str | None:
    try:
        if _docx_zip_bomb(data):
            return None
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            names = archive.namelist()
            # Reject ambiguous containers and macro-bearing files even if renamed.
            if any(name.lower().endswith("vbaproject.bin") for name in names):
                return None
            matches = [mime for marker, mime in _OFFICE_MARKERS.items() if marker in names]
            return matches[0] if len(matches) == 1 else None
    except (zipfile.BadZipFile, OSError, EOFError, NotImplementedError):
        return None


def _docx_zip_bomb(data: bytes) -> bool:
    """True when any ZIP entry or the total uncompressed size looks like a bomb.

    Per-entry cap matches ``MAX_ATTACHMENT_SIZE``; total uncompressed across
    all entries is capped at 32 MiB so a small compressed upload can't expand
    into gigabytes during DOCX parse.
    """
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            if len(archive.infolist()) > 4096:
                return True
            total = 0
            for info in archive.infolist():
                if info.file_size > MAX_ATTACHMENT_SIZE:
                    return True
                total += info.file_size
                if total > _MAX_DOCX_TOTAL_UNCOMPRESSED:
                    return True
        return False
    except (zipfile.BadZipFile, OSError, EOFError, NotImplementedError):
        return False


def _is_docx_zip(data: bytes) -> bool:
    """True only for a ZIP archive that contains the canonical DOCX marker
    entry, `word/document.xml`.

    The ZIP local-file-header signature (`PK\\x03\\x04`) alone is shared by
    .xlsx (`xl/workbook.xml`), .pptx (`ppt/presentation.xml`), .jar, .apk, and
    plain .zip files, so it isn't enough to confirm a claimed DOCX upload is
    actually a DOCX. This only inspects the archive's namelist — it never
    reads/parses `document.xml` itself.

    Runs on untrusted, potentially adversarial or truncated bytes: any
    failure to open the buffer as a well-formed ZIP is treated as "not a
    DOCX" rather than raised. Zip bombs are also rejected.
    """
    try:
        if _docx_zip_bomb(data):
            return False
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            return _DOCX_MARKER_ENTRY in archive.namelist()
    except (zipfile.BadZipFile, OSError, EOFError, NotImplementedError) as exc:
        logger.debug("ZIP signature present but not a valid/DOCX archive: %s", exc)
        return False


def _sniff_signature(data: bytes) -> str | None:
    """Detect a few common binary types by leading magic bytes."""
    if len(data) >= 3 and data[:3] == b"\xff\xd8\xff":
        return "image/jpeg"
    if len(data) >= 8 and data[:8] == b"\x89PNG\r\n\x1a\n":
        return "image/png"
    if len(data) >= 4 and data[:4] == b"%PDF":
        return "application/pdf"
    if len(data) >= 6 and data[:6] in (b"GIF87a", b"GIF89a"):
        return "image/gif"
    if len(data) >= 12 and data[:4] == b"RIFF" and data[8:12] == b"WEBP":
        return "image/webp"
    if len(data) >= 4 and data[:4] == b"PK\x03\x04":
        return _office_content_type(data)
    if len(data) >= 8 and data[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
        return "application/msword"
    return None


def _is_clean_textish(data: bytes) -> bool:
    """UTF-8 text with no NUL — rejects polyglot binaries claimed as txt/json/csv/md."""
    if b"\x00" in data:
        return False
    try:
        data.decode("utf-8")
    except UnicodeDecodeError:
        return False
    return True


def bytes_match_claimed(claimed: str, data: bytes) -> bool:
    """True if the uploaded bytes are consistent with the claimed content type.

    The presign step validates the *claim*; this checks the *actual bytes* on
    upload so a client can't presign "image/png" and then upload a shell script
    that later gets served back with `Content-Type: image/png` (a content-
    spoofing risk for any externally-shared signed URL).
    """
    norm = "image/jpeg" if claimed == "image/jpg" else claimed
    if norm in _UNVERIFIABLE_TYPES:
        return True
    detected = _sniff_signature(data)
    if detected is None:
        # No binary signature. Accept only for text-ish claims; an image/pdf
        # claim with no matching signature is a spoof.
        if norm not in _TEXTISH_TYPES:
            return False
        return _is_clean_textish(data)
    if detected != norm:
        return False
    if norm in _TEXTISH_TYPES:
        return _is_clean_textish(data)
    return True


def normalize_content_type(content_type: str) -> str:
    """Normalize client MIME types (strip params, map common aliases)."""
    base = content_type.split(";", 1)[0].strip().lower()
    return _CONTENT_TYPE_ALIASES.get(base, base)


def is_allowed_content_type(content_type: str) -> bool:
    return normalize_content_type(content_type) in ALLOWED_CONTENT_TYPES


def is_image_content_type(content_type: str) -> bool:
    return normalize_content_type(content_type) in IMAGE_CONTENT_TYPES


_DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Types extract_text_from_bytes actually knows how to parse. Used to tell a
# genuinely unsupported type (e.g. legacy .doc, which has no good pure-Python
# parser) apart from a supported type that just yielded no text (e.g. a
# scanned image-only PDF) — the two deserve different user-facing messages.
EXTRACTABLE_CONTENT_TYPES = _TEXTISH_TYPES | {"application/pdf", *_OFFICE_MARKERS.values()}


@dataclass(frozen=True)
class ExtractedText:
    text: str
    char_capped: bool = False
    page_capped: bool = False


def file_excerpt_limit_note(
    *,
    char_capped: bool,
    page_capped: bool,
    max_chars: int = MAX_EXTRACT_CHARS,
    max_pages: int = PDF_EXTRACT_MAX_PAGES,
) -> str | None:
    """Model-facing note when only a prefix of the file was read. Hidden in the UI
    because it sits inside the ``[File (type)]`` excerpt block."""
    if not char_capped and not page_capped:
        return None
    bits: list[str] = []
    if page_capped:
        bits.append(f"the first {max_pages} pages")
    if char_capped:
        bits.append(f"the first {max_chars} characters")
    covered = " and ".join(bits)
    return (
        f"[File note: only {covered} were read; later pages or text may be missing. "
        "Do not claim facts from unread parts. If asked about later content, say "
        "it was not in the excerpt.]"
    )


def _docx_tables_as_markdown(document: Any) -> list[str]:
    """Keep table cell alignment as a markdown grid instead of flattening cells."""
    blocks: list[str] = []
    for table in document.tables:
        rows_out: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            if not any(cells):
                continue
            rows_out.append("| " + " | ".join(cells) + " |")
        if not rows_out:
            continue
        col_count = max(1, rows_out[0].count("|") - 1)
        sep = "| " + " | ".join("---" for _ in range(col_count)) + " |"
        if len(rows_out) == 1:
            blocks.append(f"{rows_out[0]}\n{sep}")
        else:
            blocks.append(f"{rows_out[0]}\n{sep}\n" + "\n".join(rows_out[1:]))
    return blocks


def extract_text_details(
    content_type: str,
    data: bytes,
    *,
    max_chars: int = MAX_EXTRACT_CHARS,
) -> ExtractedText | None:
    """Sync, CPU-bound parsing with truncation flags for honest disclosure."""
    if content_type in _TEXTISH_TYPES:
        text = data.decode("utf-8", errors="replace").strip()
        if not text:
            return None
        return ExtractedText(text=text[:max_chars], char_capped=len(text) > max_chars)

    if content_type == "application/pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            page_limit = (
                PDF_INDEX_MAX_PAGES if max_chars > MAX_EXTRACT_CHARS else PDF_EXTRACT_MAX_PAGES
            )
            page_capped = len(reader.pages) > page_limit
            parts: list[str] = []
            size = 0
            char_capped = False
            for index, page in enumerate(reader.pages[:page_limit], start=1):
                page_text = page.extract_text()
                if page_text:
                    parts.append(f"[page {index}] {page_text.strip()}")
                    size += len(parts[-1]) + 2
                    if size > max_chars:
                        char_capped = True
                        break
            joined = "\n\n".join(parts).strip()
            if not joined:
                return None
            return ExtractedText(
                text=joined[:max_chars],
                char_capped=char_capped or len(joined) > max_chars,
                page_capped=page_capped,
            )
        except Exception:
            logger.debug("PDF text extraction failed", exc_info=True)
            return None

    if content_type == _DOCX_CONTENT_TYPE:
        try:
            if _docx_zip_bomb(data):
                return None
            from docx import Document

            document = Document(io.BytesIO(data))
            parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
            parts.extend(_docx_tables_as_markdown(document))
            joined = "\n".join(parts).strip()
            if not joined:
                return None
            return ExtractedText(text=joined[:max_chars], char_capped=len(joined) > max_chars)
        except Exception:
            logger.debug("DOCX text extraction failed", exc_info=True)
            return None

    if content_type in _OFFICE_MARKERS.values() and content_type != _DOCX_CONTENT_TYPE:
        if _office_content_type(data) != content_type:
            return None
        from app.services.office_text import extract_office_text

        office_text = extract_office_text(content_type, data, max_chars=max_chars + 1)
        if office_text:
            return ExtractedText(
                text=office_text[:max_chars], char_capped=len(office_text) > max_chars
            )

    # Legacy .doc (application/msword) and anything else: no parser. Handled
    # explicitly (not silently) by format_attachment_lines via
    # EXTRACTABLE_CONTENT_TYPES.
    return None


def extract_text_from_bytes(
    content_type: str,
    data: bytes,
    *,
    max_chars: int = MAX_EXTRACT_CHARS,
) -> str | None:
    """Sync, CPU-bound parsing — call via extract_text_from_bytes_async on
    any code path that isn't already off the event loop."""
    details = extract_text_details(content_type, data, max_chars=max_chars)
    return None if details is None else details.text


async def extract_text_details_async(
    content_type: str,
    data: bytes,
    settings: Settings,
    *,
    max_chars: int = MAX_EXTRACT_CHARS,
    allow_ocr: bool = True,
) -> ExtractedText | None:
    """Offload the sync, CPU-bound parse to a worker thread with a timeout,
    so a large or adversarially crafted PDF/DOCX can't block the event loop —
    same pattern as the SymPy math solve offload.

    When a PDF has no text layer, fall through to vision OCR (scanned /
    photographed pages) using a separate network timeout — unless
    *allow_ocr* is False (turn prepare; OCR belongs on the index job).
    A CPU-extract timeout does not then spend another 30s on vision.
    """
    timed_out = False
    try:
        async with asyncio.timeout(settings.attachment_extract_timeout_seconds):
            details = await asyncio.to_thread(
                extract_text_details, content_type, data, max_chars=max_chars
            )
    except TimeoutError:
        logger.warning("Attachment text extraction timed out for content_type=%s", content_type)
        timed_out = True
        details = None
    if details:
        return details
    if timed_out:
        return None
    if allow_ocr and content_type == "application/pdf" and settings.attachment_ocr_enabled:
        from app.services.attachment_ocr import ocr_scanned_pdf

        ocr_text = await ocr_scanned_pdf(settings, data)
        return ExtractedText(text=ocr_text) if ocr_text else None
    return None


async def extract_text_from_bytes_async(
    content_type: str,
    data: bytes,
    settings: Settings,
    *,
    max_chars: int = MAX_EXTRACT_CHARS,
    allow_ocr: bool = True,
) -> str | None:
    details = await extract_text_details_async(
        content_type,
        data,
        settings,
        max_chars=max_chars,
        allow_ocr=allow_ocr,
    )
    return None if details is None else details.text


async def read_attachment_bytes(gateway: StorageGateway, storage_key: str) -> bytes | None:
    return await gateway.read_bytes(storage_key)


async def verify_uploaded_bytes(
    gateway: StorageGateway,
    *,
    content_type: str,
    storage_key: str,
    max_size: int = MAX_ATTACHMENT_SIZE,
    declared_size: int | None = None,
) -> tuple[bytes | None, str | None]:
    """Return uploaded bytes when they match the declared type, else an error detail.

    When *declared_size* is provided, the actual byte count must match exactly
    — a mismatch means the client lied about the size at presign time (or the
    upload was truncated/extended), and the stored object is not what the DB
    row claims. Without this check, a presign for a 1-byte "image/png" could
    be used to store a 10 MB blob that later gets served back as a PNG.
    """
    data = await read_attachment_bytes(gateway, storage_key)
    if not data:
        return None, "Upload not found"
    if len(data) > max_size:
        return None, "Invalid upload size"
    if declared_size is not None and len(data) != declared_size:
        return None, "Uploaded size does not match the declared size"
    if not bytes_match_claimed(content_type, data):
        return None, "Uploaded bytes do not match the declared content type"
    return data, None


async def purge_invalid_upload(
    gateway: StorageGateway,
    session: AsyncSession,
    *,
    attachment_id: UUID,
    storage_key: str,
) -> None:
    from app.repositories import attachments as attachments_repo

    await gateway.delete_bytes(storage_key)
    await attachments_repo.delete_rows(session, [attachment_id])


async def ensure_verified_or_purge(
    gateway: StorageGateway,
    session: AsyncSession,
    *,
    attachment_id: UUID,
    content_type: str,
    storage_key: str,
    declared_size: int | None = None,
) -> str | None:
    """Verify R2/S3 bytes match declared type; purge row+object on failure.

    Local dev uploads are validated on PUT /upload — no-op for that backend.
    Returns an error detail when verification fails (after purge).
    """
    from app.gateways.storage_gateway import LocalStorageGateway
    from app.repositories import attachments as attachments_repo

    if isinstance(gateway, LocalStorageGateway):
        return None
    _, error = await verify_uploaded_bytes(
        gateway,
        content_type=content_type,
        storage_key=storage_key,
        declared_size=declared_size,
    )
    if error:
        await purge_invalid_upload(
            gateway,
            session,
            attachment_id=attachment_id,
            storage_key=storage_key,
        )
        return error
    await attachments_repo.mark_verified(session, attachment_id)
    return None


async def format_attachment_lines(
    gateway: StorageGateway,
    *,
    attachment_id: str,
    content_type: str,
    storage_key: str,
    size_bytes: int,
    settings: Settings,
    data: bytes | None = None,
) -> tuple[list[str], bool]:
    """Return prompt lines and whether the attachment is an image.

    When *data* is provided (e.g. bytes already loaded by turn-prep verify),
    skip a second storage download for text extraction.
    """
    if is_image_content_type(content_type):
        # API path the mobile app can fetch with auth (stored in message content).
        return [f"[Image: /attachments/{attachment_id}/file]"], True

    if data is None:
        data = await read_attachment_bytes(gateway, storage_key)
    file_ref = f"[File: /attachments/{attachment_id}/file]"
    if data:
        excerpt = await extract_text_details_async(content_type, data, settings, allow_ocr=False)
        if excerpt:
            note = file_excerpt_limit_note(
                char_capped=excerpt.char_capped,
                page_capped=excerpt.page_capped,
            )
            body = f"{note}\n{excerpt.text}" if note else excerpt.text
            return [file_ref, f"[File ({content_type})]\n{body}"], False
        if content_type not in EXTRACTABLE_CONTENT_TYPES:
            return (
                [
                    file_ref,
                    f"[File attached: {content_type}. Recall can't read this file type yet — "
                    "tell the user to try a PDF, Word (.docx), or plain text file instead.]",
                ],
                False,
            )
        if content_type == "application/pdf":
            return (
                [
                    file_ref,
                    "[File attached: application/pdf. No extractable text layer — this is likely a "
                    "scanned or image-only PDF. OCR runs in the background and will be available "
                    "on later turns; suggest a text-based PDF or pasting the text if needed.]",
                ],
                False,
            )
        return (
            [
                file_ref,
                f"[File attached: {content_type}. No extractable text was found — "
                "tell the user the file appears empty or unreadable.]",
            ],
            False,
        )

    return [file_ref, f"[File attached: {content_type}, {size_bytes} bytes]"], False


def _is_image_marker_line(line: str) -> bool:
    s = line.strip()
    return s.startswith("[Image:") and s.endswith("]") and len(s) > len("[Image:]")


def _is_file_ref_line(line: str) -> bool:
    s = line.strip()
    return s.startswith("[File:") and s.endswith("]") and len(s) > len("[File:]")


def image_attachment_ids_from_text(content: str) -> list[UUID]:
    """Parse ``[Image: /attachments/{uuid}/file]`` markers with a linear scan."""
    found: list[UUID] = []
    seen: set[UUID] = set()
    for line in content.splitlines():
        stripped = line.strip()
        if not _is_image_marker_line(stripped):
            continue
        inner = stripped[len("[Image:") : -1].strip()
        start = inner.find(_ATTACH_ID_PREFIX)
        if start < 0:
            continue
        id_start = start + len(_ATTACH_ID_PREFIX)
        id_end = inner.find(_ATTACH_ID_SUFFIX, id_start)
        if id_end < 0:
            continue
        raw = inner[id_start:id_end]
        try:
            uid = UUID(raw)
        except ValueError:
            continue
        if uid in seen:
            continue
        seen.add(uid)
        found.append(uid)
    return found


def _content_as_text(content: Any) -> str:
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts: list[str] = []
        for part in content:
            if isinstance(part, dict) and part.get("type") == "text":
                text = part.get("text")
                if isinstance(text, str):
                    parts.append(text)
        return "\n".join(parts)
    return ""


def prior_image_attachment_ids(
    messages: list[dict[str, Any]],
    *,
    limit: int = VISION_REHYDRATE_LIMIT,
) -> list[UUID]:
    """Most recent prior-turn image ids (newest first), excluding the last user turn."""
    if limit <= 0:
        return []
    last_user: int | None = None
    for index in range(len(messages) - 1, -1, -1):
        if messages[index].get("role") == "user":
            last_user = index
            break
    start = last_user - 1 if last_user is not None else len(messages) - 1
    ordered: list[UUID] = []
    seen: set[UUID] = set()
    for index in range(start, -1, -1):
        for uid in image_attachment_ids_from_text(_content_as_text(messages[index].get("content"))):
            if uid in seen:
                continue
            seen.add(uid)
            ordered.append(uid)
            if len(ordered) >= limit:
                return ordered
    return ordered


def history_has_image_marker(messages: list[Any]) -> bool:
    """True when any history row still has an ``[Image: …]`` marker."""
    for row in messages:
        if isinstance(row, dict):
            text = _content_as_text(row.get("content"))
        else:
            raw = getattr(row, "content", None)
            text = raw if isinstance(raw, str) else ""
        if image_attachment_ids_from_text(text):
            return True
    return False


def append_image_unavailable_note(prompt_messages: list[dict[str, Any]]) -> None:
    """Tell the model it cannot see a prior image — do not guess from a caption."""
    for index in range(len(prompt_messages) - 1, -1, -1):
        msg = prompt_messages[index]
        if msg.get("role") != "user":
            continue
        content = msg.get("content")
        if isinstance(content, str):
            msg["content"] = f"{content}\n\n{_IMAGE_UNAVAILABLE_NOTE}"
            return
        if isinstance(content, list):
            content.append({"type": "text", "text": _IMAGE_UNAVAILABLE_NOTE})
            return
        msg["content"] = _IMAGE_UNAVAILABLE_NOTE
        return


def strip_attachment_from_content(content: str, attachment_id: UUID) -> str:
    """Remove this attachment's markers (and a following file excerpt) from a message.

    Linear scan — no regex. Empty result means the caller should delete the row.
    """
    needle = str(attachment_id)
    kept: list[str] = []
    skipping_file_block = False
    for line in content.split("\n"):
        stripped = line.strip()
        if skipping_file_block:
            if _is_image_marker_line(stripped) or _is_file_ref_line(stripped):
                skipping_file_block = False
            else:
                continue
        if needle in stripped and (_is_image_marker_line(stripped) or _is_file_ref_line(stripped)):
            skipping_file_block = _is_file_ref_line(stripped)
            continue
        kept.append(line)
    return "\n".join(kept).strip()


def _strip_image_markers(text: str) -> str:
    lines = [line for line in text.splitlines() if not _is_image_marker_line(line)]
    return "\n".join(lines).strip()


async def inject_vision_content(
    prompt_messages: list[dict[str, Any]],
    gateway: StorageGateway,
    images: list[tuple[str, str]],
    *,
    caption: str = "",
    bytes_by_key: dict[str, bytes] | None = None,
) -> bool:
    """Replace the last user turn with multimodal content for vision models.

    Prefer *bytes_by_key* (turn-prep verify cache) over re-downloading.
    Returns True when at least one image part was injected.
    """
    parts: list[dict[str, Any]] = []
    text = _strip_image_markers(caption).strip() or "What's in this image?"
    parts.append({"type": "text", "text": text})

    async def _bytes_for(storage_key: str) -> bytes | None:
        if bytes_by_key is not None and storage_key in bytes_by_key:
            return bytes_by_key[storage_key]
        return await read_attachment_bytes(gateway, storage_key)

    image_bytes = await asyncio.gather(
        *(_bytes_for(storage_key) for _content_type, storage_key in images)
    )
    for (content_type, _storage_key), data in zip(images, image_bytes, strict=True):
        if not data:
            continue
        mime = normalize_content_type(content_type)
        encoded = base64.standard_b64encode(data).decode("ascii")
        parts.append(
            {
                "type": "image_url",
                "image_url": {"url": f"data:{mime};base64,{encoded}"},
            }
        )

    if not any(part.get("type") == "image_url" for part in parts):
        return False

    for idx in range(len(prompt_messages) - 1, -1, -1):
        if prompt_messages[idx].get("role") == "user":
            prompt_messages[idx] = {"role": "user", "content": parts}
            return True
    return False
